import sqlite3
import random
import pandas as pd
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

DB_PATH = Path("database/question_bank.db")
OUTPUT_DIR = Path("outputs")
OUTPUT_DIR.mkdir(exist_ok=True)

def buscar_questoes(termo, quantidade):
    conn = sqlite3.connect(DB_PATH)

    query = """
    SELECT
        qe.id,
        qe.numero_questao,
        qe.texto_questao,
        qe.area,
        qe.subarea,
        qe.tema,
        qe.competencia,
        qe.periodo,
        d.prova,
        d.instituicao,
        d.ano
    FROM questoes_extraidas qe
    LEFT JOIN documentos d ON qe.documento_id = d.id
    WHERE qe.qualidade = 'valida' AND
        qe.area LIKE ?
        OR qe.subarea LIKE ?
        OR qe.tema LIKE ?
        OR qe.competencia LIKE ?
        OR qe.texto_questao LIKE ?
    """

    param = f"%{termo}%"
    df = pd.read_sql_query(query, conn, params=[param, param, param, param, param])
    conn.close()

    if len(df) == 0:
        return df

    quantidade = min(quantidade, len(df))
    return df.sample(n=quantidade, random_state=random.randint(1, 999999))

def limpar_texto(texto):
    texto = str(texto)
    texto = texto.replace("\uFFFE", "")
    texto = texto.replace("￾", "-")
    texto = texto.replace("\r", "")
    return texto.strip()

def gerar_word(df, termo):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("ATHENA QUESTION BANK")
    run.bold = True
    run.font.size = Pt(16)

    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitulo.add_run("PROVA GERADA AUTOMATICAMENTE")
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph(f"Critério de busca: {termo}")
    doc.add_paragraph(f"Número de questões: {len(df)}")
    doc.add_paragraph("")

    for i, row in enumerate(df.itertuples(), start=1):
        p = doc.add_paragraph()
        run = p.add_run(f"QUESTÃO {i}")
        run.bold = True
        run.font.size = Pt(12)

        texto = limpar_texto(row.texto_questao)
        p = doc.add_paragraph(texto)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        fonte = (
            f"Fonte: {row.prova} | Instituição: {row.instituicao} | "
            f"Ano: {row.ano} | Área: {row.area} | Tema: {row.tema}"
        )

        p = doc.add_paragraph()
        run = p.add_run(fonte)
        run.italic = True
        run.font.size = Pt(9)

        doc.add_paragraph("_" * 80)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S"); nome = f"prova_{termo.replace(' ', '_')}_{timestamp}.docx"
    caminho = OUTPUT_DIR / nome
    doc.save(caminho)
    return caminho

def gerar_excel(df, termo):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S"); nome = f"prova_{termo.replace(' ', '_')}_{timestamp}.xlsx"
    caminho = OUTPUT_DIR / nome
    df.to_excel(caminho, index=False)
    return caminho

def main():
    print("=" * 60)
    print("GERADOR DE PROVAS - ATHENA QUESTION BANK")
    print("=" * 60)

    termo = input("Digite a área, tema ou palavra-chave: ").strip()
    quantidade = int(input("Número de questões desejado: ").strip())

    df = buscar_questoes(termo, quantidade)

    if len(df) == 0:
        print("Nenhuma questão encontrada.")
        return

    caminho_word = gerar_word(df, termo)
    caminho_excel = gerar_excel(df, termo)

    print("\nProva gerada com sucesso.")
    print(f"Word: {caminho_word}")
    print(f"Excel: {caminho_excel}")

if __name__ == "__main__":
    main()
