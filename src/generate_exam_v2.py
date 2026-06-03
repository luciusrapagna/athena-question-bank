import sqlite3
import random
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt

DB_PATH = Path("database/question_bank.db")
OUTPUT_DIR = Path("outputs")
OUTPUT_DIR.mkdir(exist_ok=True)

def buscar(termo, quantidade):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()

    param = f"%{termo}%"

    cur.execute("""
    SELECT
        qs.numero_questao,
        qs.enunciado,
        qs.alternativa_a,
        qs.alternativa_b,
        qs.alternativa_c,
        qs.alternativa_d,
        qe.area,
        qe.subarea,
        qe.tema,
        d.prova,
        d.instituicao,
        d.ano
    FROM questoes_struct qs
    LEFT JOIN questoes_extraidas qe
        ON qs.documento_id = qe.documento_id
       AND qs.numero_questao = qe.numero_questao
    LEFT JOIN documentos d
        ON qs.documento_id = d.id
    WHERE qs.qualidade = 'valida_struct'
      AND (
        qs.enunciado LIKE ?
        OR qs.alternativa_a LIKE ?
        OR qs.alternativa_b LIKE ?
        OR qs.alternativa_c LIKE ?
        OR qs.alternativa_d LIKE ?
        OR qe.area LIKE ?
        OR qe.subarea LIKE ?
        OR qe.tema LIKE ?
      )
    """, (param, param, param, param, param, param, param, param))

    rows = cur.fetchall()
    conn.close()

    random.shuffle(rows)
    return rows[:quantidade]

def gerar_word(questoes, termo):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    doc.add_heading("ATHENA QUESTION BANK", level=1)
    doc.add_heading("PROVA GERADA AUTOMATICAMENTE", level=2)

    doc.add_paragraph(f"Critério de busca: {termo}")
    doc.add_paragraph(f"Número de questões: {len(questoes)}")

    for i, q in enumerate(questoes, start=1):
        numero, enunciado, a, b, c, d, area, subarea, tema, prova, instituicao, ano = q

        doc.add_paragraph("")
        p = doc.add_paragraph()
        r = p.add_run(f"QUESTÃO {i}")
        r.bold = True

        doc.add_paragraph(enunciado)

        doc.add_paragraph(f"(A) {a}")
        doc.add_paragraph(f"(B) {b}")
        doc.add_paragraph(f"(C) {c}")
        doc.add_paragraph(f"(D) {d}")

        fonte = doc.add_paragraph()
        run = fonte.add_run(
            f"Fonte: {prova} | Instituição: {instituicao} | Ano: {ano} | "
            f"Área: {area} | Subárea: {subarea} | Tema: {tema} | Questão original: {numero}"
        )
        run.italic = True
        run.font.size = Pt(9)

        doc.add_paragraph("_" * 80)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    caminho = OUTPUT_DIR / f"prova_v2_{termo.replace(' ', '_')}_{timestamp}.docx"
    doc.save(caminho)

    return caminho

def main():
    print("=" * 60)
    print("ATHENA EXAM GENERATOR V2")
    print("=" * 60)

    termo = input("Digite área, tema ou palavra-chave: ").strip()
    quantidade = int(input("Número de questões: ").strip())

    questoes = buscar(termo, quantidade)

    if not questoes:
        print("Nenhuma questão encontrada.")
        return

    caminho = gerar_word(questoes, termo)

    print(f"Prova gerada: {caminho}")

if __name__ == "__main__":
    main()
