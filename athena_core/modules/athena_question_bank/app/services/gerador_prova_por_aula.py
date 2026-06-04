import sqlite3
from pathlib import Path
from docx import Document

DB_PATH = "app/db/planos_aula.db"
OUTPUT_DIR = Path("outputs/provas_por_aula")


def criar_pasta_saida():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def buscar_planos(cur):
    return cur.execute("""
        SELECT id, aula
        FROM planos_aula
        ORDER BY id
    """).fetchall()


def buscar_questoes_compativeis(cur, plano_id, limite=10):
    return cur.execute("""
        SELECT c.questao_id, ROUND(c.score, 3)
        FROM compatibilidade_plano_questao c
        WHERE c.plano_id = ?
        ORDER BY c.score DESC
        LIMIT ?
    """, (plano_id, limite)).fetchall()


def buscar_texto_questao(cur, questao_id):
    colunas = [c[1] for c in cur.execute("PRAGMA table_info(questoes);").fetchall()]

    for coluna in ["enunciado", "texto", "questao", "conteudo", "pergunta"]:
        if coluna in colunas:
            row = cur.execute(
                f"SELECT {coluna} FROM questoes WHERE id = ?",
                (questao_id,)
            ).fetchone()
            return row[0] if row else ""

    return ""


def gerar_docx(plano_id, aula, questoes):
    doc = Document()

    doc.add_heading("ATHENA QUESTION BANK", level=0)
    doc.add_heading("Prova Gerada Automaticamente por Aula", level=1)

    doc.add_paragraph(f"Plano/Aula: {aula}")
    doc.add_paragraph(f"ID do plano: {plano_id}")
    doc.add_paragraph(f"Total de questões: {len(questoes)}")

    doc.add_heading("Questões selecionadas", level=2)

    for i, item in enumerate(questoes, start=1):
        questao_id, score, texto = item
        doc.add_paragraph(f"Questão {i} — ID {questao_id} — Compatibilidade: {score}")
        doc.add_paragraph(texto)
        doc.add_paragraph("")

    nome_seguro = "".join(c for c in aula if c.isalnum() or c in " _-").strip()
    if not nome_seguro:
        nome_seguro = f"plano_{plano_id}"

    caminho = OUTPUT_DIR / f"prova_aula_{plano_id}_{nome_seguro}.docx"
    doc.save(caminho)

    return caminho


def registrar_prova(cur, plano_id, total, observacoes):
    cur.execute("""
        INSERT INTO provas_por_aula (
            plano_id,
            total_questoes,
            observacoes
        )
        VALUES (?, ?, ?)
    """, (plano_id, total, observacoes))


def gerar_provas_por_aula(limite_questoes=10):
    criar_pasta_saida()

    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()

    planos = buscar_planos(cur)

    total_provas = 0

    for plano_id, aula in planos:
        compat = buscar_questoes_compativeis(cur, plano_id, limite_questoes)

        if not compat:
            print(f"Plano {plano_id} sem questões compatíveis.")
            continue

        questoes = []

        for questao_id, score in compat:
            texto = buscar_texto_questao(cur, questao_id)
            questoes.append((questao_id, score, texto))

        caminho = gerar_docx(plano_id, aula or f"Plano {plano_id}", questoes)

        registrar_prova(
            cur,
            plano_id,
            len(questoes),
            f"Prova gerada automaticamente: {caminho}"
        )

        print(f"Prova gerada: {caminho}")
        total_provas += 1

    con.commit()
    con.close()

    print(f"Total de provas geradas: {total_provas}")


if __name__ == "__main__":
    gerar_provas_por_aula()
