import csv
import argparse
from pathlib import Path
from docx import Document
from app.config_paths import BANCO_QUESTOES, BANCO_PLANOS, BANCO_RELACIONAMENTOS, DIR_WORD


def carregar_csv(caminho):
    with open(caminho, "r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def selecionar_questoes(aula, area, quantidade, assunto="", competencia=""):
    questoes = {q["id_questao"]: q for q in carregar_csv(BANCO_QUESTOES)}
    planos = {p["id_aula"]: p for p in carregar_csv(BANCO_PLANOS)}
    relacoes = carregar_csv(BANCO_RELACIONAMENTOS)

    aula = aula.lower().strip()
    area = area.lower().strip()

    ids_aulas = [
        p["id_aula"] for p in planos.values()
        if aula in p["aula"].lower() or aula in p["assunto"].lower()
    ]

    selecionadas = []

    for r in sorted(relacoes, key=lambda x: float(x["score_aderencia"]), reverse=True):
        if r["id_aula"] not in ids_aulas:
            continue

        q = questoes.get(r["id_questao"])
        if not q:
            continue

        if area and q["area"].lower() != area:
            continue

        if assunto and assunto not in q.get("assunto", "").lower():
            continue

        if competencia and competencia not in q.get("competencia", "").lower():
            continue

        q["_score_aderencia"] = r["score_aderencia"]
        selecionadas.append(q)

        if len(selecionadas) >= quantidade:
            break

    return selecionadas


def exportar_word(questoes, aula, area, quantidade, assunto="", competencia=""):
    DIR_WORD.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading("ATHENA QUESTION BANK", level=1)
    doc.add_heading("Questões selecionadas por plano de aula", level=2)

    doc.add_paragraph(f"Aula solicitada: {aula}")
    doc.add_paragraph(f"Grande área: {area}")
    doc.add_paragraph(f"Assunto: {assunto or ''}")
    doc.add_paragraph(f"Competência: {competencia or ''}")
    doc.add_paragraph(f"Quantidade solicitada: {quantidade}")
    doc.add_paragraph(f"Questões encontradas: {len(questoes)}")

    for i, q in enumerate(questoes, start=1):
        doc.add_heading(f"Questão {i}", level=2)
        doc.add_paragraph(f"Fonte: {q.get('fonte', '')} | Ano: {q.get('ano', '')} | Área: {q.get('area', '')}")
        doc.add_paragraph(f"Score de aderência: {q.get('_score_aderencia', '')}")
        doc.add_paragraph(q.get("enunciado", ""))

        for letra, campo in [
            ("A", "alternativa_a"),
            ("B", "alternativa_b"),
            ("C", "alternativa_c"),
            ("D", "alternativa_d"),
            ("E", "alternativa_e"),
        ]:
            alt = q.get(campo, "")
            if alt:
                doc.add_paragraph(f"{letra}) {alt}")

        if q.get("gabarito"):
            doc.add_paragraph(f"Gabarito: {q.get('gabarito')}")

    nome = f"questoes_{aula}_{area}_{assunto}_{competencia}_{quantidade}.docx"
    nome = nome.lower().replace(" ", "_").replace("ç", "c").replace("ú", "u").replace("ã", "a")
    caminho = DIR_WORD / nome
    doc.save(caminho)

    return caminho


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--aula", default="")
    parser.add_argument("--area", default="")
    parser.add_argument("--assunto", default="")
    parser.add_argument("--competencia", default="")
    parser.add_argument("--n", type=int, default=3)

    args = parser.parse_args()

    questoes = selecionar_questoes(args.aula, args.area, args.n, args.assunto.lower().strip(), args.competencia.lower().strip())
    caminho = exportar_word(questoes, args.aula, args.area, args.n, args.assunto, args.competencia)

    print(f"Questões selecionadas: {len(questoes)}")
    print(f"Word gerado em: {caminho}")


if __name__ == "__main__":
    main()




