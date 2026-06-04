import csv
import argparse
from docx import Document
from app.config_paths import BANCO_QUESTOES, DIR_SIMULADOS


def carregar_questoes():
    with open(BANCO_QUESTOES, "r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def filtrar_questoes(area, assunto=None):
    questoes = carregar_questoes()
    area = area.lower().strip()
    assunto = assunto.lower().strip() if assunto else None

    filtradas = []

    for q in questoes:
        texto = " ".join([
            q.get("area", ""),
            q.get("assunto", ""),
            q.get("enunciado", "")
        ]).lower()

        if q.get("area", "").lower().strip() != area:
            continue

        if assunto and assunto not in texto:
            continue

        filtradas.append(q)

    return filtradas


def exportar_simulado(configs, assunto=None, nome_saida="simulado_personalizado.docx"):
    DIR_SIMULADOS.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading("ATHENA QUESTION BANK", level=1)
    doc.add_heading("Simulado personalizado", level=2)

    if assunto:
        doc.add_paragraph(f"Assunto direcionador: {assunto}")

    total = 0

    for area, quantidade in configs.items():
        questoes = filtrar_questoes(area, assunto=assunto)
        selecionadas = questoes[:quantidade]

        doc.add_heading(area, level=2)
        doc.add_paragraph(f"Questões solicitadas: {quantidade}")
        doc.add_paragraph(f"Questões encontradas: {len(selecionadas)}")

        for q in selecionadas:
            total += 1
            doc.add_heading(f"Questão {total}", level=3)
            doc.add_paragraph(f"Fonte: {q.get('fonte', '')} | Ano: {q.get('ano', '')} | Área: {q.get('area', '')}")
            doc.add_paragraph(q.get("enunciado", ""))

            alternativas = [
                ("A", "alternativa_a"),
                ("B", "alternativa_b"),
                ("C", "alternativa_c"),
                ("D", "alternativa_d"),
                ("E", "alternativa_e"),
            ]

            for letra, campo in alternativas:
                alt = q.get(campo, "")
                if alt:
                    doc.add_paragraph(f"{letra}) {alt}")

            if q.get("gabarito"):
                doc.add_paragraph(f"Gabarito: {q.get('gabarito')}")

    caminho = DIR_SIMULADOS / nome_saida
    doc.save(caminho)

    print(f"Simulado gerado: {caminho}")
    print(f"Total de questões no simulado: {total}")


def main():
    parser = argparse.ArgumentParser()

    parser.add_argument("--clinica", type=int, default=0)
    parser.add_argument("--cirurgia", type=int, default=0)
    parser.add_argument("--pediatria", type=int, default=0)
    parser.add_argument("--go", type=int, default=0)
    parser.add_argument("--saude_coletiva", type=int, default=0)
    parser.add_argument("--assunto", default=None)

    args = parser.parse_args()

    configs = {
        "Clínica Médica": args.clinica,
        "Cirurgia": args.cirurgia,
        "Pediatria": args.pediatria,
        "Ginecologia e Obstetrícia": args.go,
        "Saúde Coletiva": args.saude_coletiva,
    }

    configs = {area: n for area, n in configs.items() if n > 0}

    exportar_simulado(
        configs=configs,
        assunto=args.assunto,
    )


if __name__ == "__main__":
    main()
