import csv
from pathlib import Path
from docx import Document
from app.config_paths import BANCO_QUESTOES

BASE_DIR = Path(__file__).resolve().parents[2]
DIR_SIMULADOS = BASE_DIR / "outputs" / "simulados"
DIR_SIMULADOS.mkdir(parents=True, exist_ok=True)

GRANDES_AREAS = [
    "Clínica Médica",
    "Cirurgia",
    "Pediatria",
    "Ginecologia e Obstetrícia",
    "Saúde Coletiva",
]


def carregar_questoes():
    with open(BANCO_QUESTOES, "r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def escolher_assuntos(area, questoes):
    assuntos = sorted({
        q["assunto"] for q in questoes
        if q.get("area") == area and q.get("assunto")
    })

    print(f"\n{area}")
    print("-" * 60)

    for i, assunto in enumerate(assuntos, start=1):
        total = sum(1 for q in questoes if q.get("area") == area and q.get("assunto") == assunto)
        print(f"[ ] {i} - {assunto} ({total} questões)")

    escolha = input("\nDigite os números dos assuntos desejados separados por vírgula ou ENTER para todos: ").strip()

    if not escolha:
        return assuntos

    indices = []
    for item in escolha.split(","):
        item = item.strip()
        if item.isdigit():
            indices.append(int(item))

    return [
        assuntos[i - 1]
        for i in indices
        if 1 <= i <= len(assuntos)
    ]


def gerar_simulado():
    questoes = carregar_questoes()
    selecionadas = []

    print("ATHENA SIMULADO GENERATOR")
    print("=" * 60)

    for area in GRANDES_AREAS:
        total_area = sum(1 for q in questoes if q.get("area") == area)

        print(f"\nÁrea: {area} | Disponíveis: {total_area}")
        usar = input("Deseja incluir esta área? [s/n]: ").strip().lower()

        if usar != "s":
            continue

        assuntos_escolhidos = escolher_assuntos(area, questoes)

        qtd = input(f"Quantas questões deseja incluir de {area}? ").strip()

        if not qtd.isdigit():
            print("Quantidade inválida. Área ignorada.")
            continue

        qtd = int(qtd)

        filtradas = [
            q for q in questoes
            if q.get("area") == area and q.get("assunto") in assuntos_escolhidos
        ]

        selecionadas.extend(filtradas[:qtd])

    exportar_word(selecionadas)


def exportar_word(questoes):
    doc = Document()

    doc.add_heading("ATHENA QUESTION BANK", level=1)
    doc.add_heading("Simulado Personalizado", level=2)

    doc.add_paragraph(f"Total de questões: {len(questoes)}")

    for i, q in enumerate(questoes, start=1):
        doc.add_heading(f"Questão {i}", level=2)
        doc.add_paragraph(f"Área: {q.get('area', '')}")
        doc.add_paragraph(f"Assunto: {q.get('assunto', '')}")
        doc.add_paragraph(f"Competência: {q.get('competencia', '')}")
        doc.add_paragraph(q.get("enunciado", ""))

        for letra, campo in [
            ("A", "alternativa_a"),
            ("B", "alternativa_b"),
            ("C", "alternativa_c"),
            ("D", "alternativa_d"),
            ("E", "alternativa_e"),
        ]:
            if q.get(campo):
                doc.add_paragraph(f"{letra}) {q.get(campo)}")

    caminho = DIR_SIMULADOS / "simulado_personalizado_interativo.docx"
    doc.save(caminho)

    print("\nSimulado gerado com sucesso.")
    print(f"Arquivo: {caminho}")


if __name__ == "__main__":
    gerar_simulado()
