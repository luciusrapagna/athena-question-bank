import sqlite3
import re

DB = "app/db/planos_aula.db"

FALLBACK = "Tema médico não classificado"

def normalizar(txt):
    if not txt:
        return ""
    return re.sub(r"\s+", " ", str(txt).lower()).strip()

def limpar(txt):
    if not txt or str(txt).strip() == "" or txt == "Assunto não identificado":
        return FALLBACK
    return str(txt).strip()

def buscar_opcoes(campo):
    con = sqlite3.connect(DB)
    cur = con.cursor()
    cur.execute(f"""
        SELECT DISTINCT {campo}
        FROM questoes
        WHERE {campo} IS NOT NULL
          AND TRIM(CAST({campo} AS TEXT)) <> ''
          AND {campo} <> 'Assunto não identificado'
        ORDER BY {campo}
    """)
    dados = [r[0] for r in cur.fetchall()]
    con.close()
    return dados

def selecionar_questoes_inteligente(
    area=None,
    especialidade=None,
    tema=None,
    subtema=None,
    competencia=None,
    limite=10
):
    con = sqlite3.connect(DB)
    con.row_factory = sqlite3.Row
    cur = con.cursor()

    cur.execute("SELECT * FROM questoes")
    questoes = [dict(r) for r in cur.fetchall()]
    con.close()

    resultados = []

    for q in questoes:
        score = 0

        if area and normalizar(area) in normalizar(q.get("area") or q.get("grande_area") or q.get("categoria")):
            score += 20

        if especialidade and normalizar(especialidade) == normalizar(q.get("especialidade")):
            score += 30

        if tema and (
            normalizar(tema) == normalizar(q.get("tema_indexado"))
            or normalizar(tema) in normalizar(q.get("assunto"))
        ):
            score += 40

        if subtema and normalizar(subtema) == normalizar(q.get("subtema_indexado")):
            score += 35

        if competencia and normalizar(competencia) == normalizar(q.get("competencia_enamed")):
            score += 25

        sem_filtros = not any([area, especialidade, tema, subtema, competencia])

        if score > 0 or sem_filtros:
            q["_score_selector"] = score
            resultados.append(q)

    resultados = sorted(
        resultados,
        key=lambda x: (
            x.get("_score_selector", 0),
            x.get("confianca_indexacao") or 0
        ),
        reverse=True
    )

    return resultados[:int(limite)]

def texto_enunciado(q):
    return q.get("enunciado") or q.get("texto") or q.get("questao") or ""

def exportar_word(questoes, caminho):
    from docx import Document

    doc = Document()
    doc.add_heading("ATHENA Question Bank", level=0)
    doc.add_heading("Question Selector Inteligente — Sprint 8", level=1)

    for i, q in enumerate(questoes, 1):
        doc.add_heading(f"Questão {i}", level=2)

        campos = [
            ("Área", q.get("area") or q.get("grande_area") or q.get("categoria")),
            ("Especialidade", q.get("especialidade")),
            ("Tema", q.get("tema_indexado") or q.get("assunto")),
            ("Subtema", q.get("subtema_indexado")),
        ]

        for nome, valor in campos:
            p = doc.add_paragraph()
            p.add_run(f"{nome}: ").bold = True
            p.add_run(limpar(valor))

        doc.add_paragraph(texto_enunciado(q))

        for letra in ["a", "b", "c", "d", "e"]:
            valor = q.get(f"alternativa_{letra}")
            if valor:
                doc.add_paragraph(f"{letra.upper()}) {valor}")

    doc.save(caminho)
    return caminho
