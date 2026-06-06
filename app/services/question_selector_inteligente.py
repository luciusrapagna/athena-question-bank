import sqlite3
import re

DB = "app/db/planos_aula.db"
FALLBACK = "Não classificado"

def normalizar(txt):
    if not txt:
        return ""
    return re.sub(r"\s+", " ", str(txt).lower()).strip()

def limpar(txt):
    if not txt or str(txt).strip() == "" or txt == "Assunto não identificado":
        return FALLBACK
    return str(txt).strip()

def buscar_opcoes(campo):
    con = sqlite3.connect(DB)
    cur = con.cursor()
    cur.execute(f"""
        SELECT DISTINCT {campo}
        FROM questoes
        WHERE {campo} IS NOT NULL
          AND TRIM(CAST({campo} AS TEXT)) <> ''
          AND {campo} <> 'Assunto não identificado'
        ORDER BY {campo}
    """)
    dados = [r[0] for r in cur.fetchall()]
    con.close()
    return dados

def selecionar_questoes_inteligente(
    area=None,
    especialidade=None,
    tema=None,
    subtema=None,
    competencia=None,
    limite=10
):
    con = sqlite3.connect(DB)
    con.row_factory = sqlite3.Row
    cur = con.cursor()

    cur.execute("SELECT * FROM questoes")
    questoes = [dict(r) for r in cur.fetchall()]
    con.close()

    resultados = []

    for q in questoes:
        score = 0

        area_q = q.get("grande_area") or q.get("area") or ""
        tema_q = q.get("tema_indexado") or q.get("assunto") or q.get("tema") or ""
        subtema_q = q.get("subtema_indexado") or ""
        esp_q = q.get("especialidade") or q.get("especialidade_indexada") or ""
        comp_q = q.get("competencia_enamed") or q.get("competencia") or ""

        if area and normalizar(area) in normalizar(area_q):
            score += 20

        if especialidade and normalizar(especialidade) == normalizar(esp_q):
            score += 30

        if tema and (
            normalizar(tema) == normalizar(tema_q)
            or normalizar(tema) in normalizar(q.get("enunciado"))
        ):
            score += 40

        if subtema and normalizar(subtema) == normalizar(subtema_q):
            score += 35

        if competencia and normalizar(competencia) == normalizar(comp_q):
            score += 25

        sem_filtros = not any([area, especialidade, tema, subtema, competencia])

        if score > 0 or sem_filtros:
            q["_score_selector"] = score
            resultados.append(q)

    resultados = sorted(
        resultados,
        key=lambda x: (
            x.get("_score_selector", 0),
            x.get("confianca_indexacao") or 0
        ),
        reverse=True
    )

    return resultados[:int(limite)]

def texto_enunciado(q):
    return q.get("enunciado") or q.get("texto") or q.get("questao") or ""

def exportar_word(questoes, caminho):
    from docx import Document

    doc = Document()
    doc.add_heading("ATHENA Question Bank", level=0)
    doc.add_heading("Question Selector por Semana Inteligente — Sprint 9", level=1)

    for i, q in enumerate(questoes, 1):
        doc.add_heading(f"Questão {i}", level=2)

        campos = [
            ("Tema da Aula", q.get("tema_aula") or q.get("conteudo_semana_area")),
            ("Área", q.get("grande_area") or q.get("area")),
            ("Especialidade", q.get("especialidade") or q.get("especialidade_indexada")),
            ("Tema Indexado", q.get("tema_indexado") or q.get("assunto") or q.get("tema")),
            ("Subtema", q.get("subtema_indexado")),
            ("Competência ENAMED", q.get("competencia_enamed") or q.get("competencia")),
            ("Habilidade ENAMED", q.get("habilidade_enamed") or q.get("habilidade")),
        ]

        for nome, valor in campos:
            p = doc.add_paragraph()
            p.add_run(f"{nome}: ").bold = True
            p.add_run(limpar(valor))

        doc.add_paragraph(texto_enunciado(q))

        for letra in ["a", "b", "c", "d", "e"]:
            valor = q.get(f"alternativa_{letra}")
            if valor:
                doc.add_paragraph(f"{letra.upper()}) {valor}")

        gab = q.get("gabarito")
        if gab:
            p = doc.add_paragraph()
            p.add_run("Gabarito: ").bold = True
            p.add_run(str(gab))

    doc.save(caminho)
    return caminho
