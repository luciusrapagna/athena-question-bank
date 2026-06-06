import sqlite3
import re
from pathlib import Path
from datetime import datetime
from docx import Document

DB = "app/db/planos_aula.db"
OUTDIR = Path("outputs/word")
OUTDIR.mkdir(parents=True, exist_ok=True)

FALLBACK = "Tema médico não classificado"

def normalizar(txt):
    if not txt:
        return ""
    txt = str(txt).lower().strip()
    txt = re.sub(r"\s+", " ", txt)
    return txt

def limpar(txt):
    if not txt:
        return FALLBACK
    txt = str(txt).strip()
    if txt == "" or txt == "Assunto não identificado":
        return FALLBACK
    return txt

def similaridade(a, b):
    a = set(normalizar(a).split())
    b = set(normalizar(b).split())
    if not a or not b:
        return 0
    return len(a & b) / len(a | b)

def calcular_score(aula, questao):
    tema_aula = aula.get("tema") or aula.get("assunto") or ""
    periodo_aula = aula.get("periodo") or aula.get("semana") or ""

    tema_q = questao.get("tema_indexado") or questao.get("assunto") or ""
    subtema_q = questao.get("subtema_indexado") or ""
    esp_q = questao.get("especialidade") or ""
    comp_q = questao.get("competencia_enamed") or ""
    periodo_q = questao.get("periodo") or ""

    score_tema = max(
        similaridade(tema_aula, tema_q),
        similaridade(tema_aula, subtema_q),
        similaridade(tema_aula, comp_q),
    )

    score_especialidade = similaridade(tema_aula, esp_q)

    score_periodo = 0
    if periodo_aula and periodo_q:
        score_periodo = 1 if normalizar(periodo_aula) == normalizar(periodo_q) else 0.3
    else:
        score_periodo = 0.5

    return round((score_tema * 0.70) + (score_especialidade * 0.20) + (score_periodo * 0.10), 4)

def carregar_aulas(cur, aula_numero=None, semana=None):
    cur.execute("PRAGMA table_info(planos_aula)")
    cols = [c[1] for c in cur.fetchall()]

    where = []
    params = []

    if aula_numero is not None and "aula" in cols:
        where.append("aula = ?")
        params.append(aula_numero)

    if semana is not None and "semana" in cols:
        where.append("semana = ?")
        params.append(semana)

    sql = "SELECT * FROM planos_aula"
    if where:
        sql += " WHERE " + " AND ".join(where)

    cur.execute(sql, params)
    return [dict(r) for r in cur.fetchall()]

def carregar_questoes(cur):
    cur.execute("SELECT * FROM questoes")
    return [dict(r) for r in cur.fetchall()]

def selecionar_questoes(aulas, questoes, n_por_area=2):
    selecionadas = []

    for aula in aulas:
        area_aula = aula.get("area") or aula.get("grande_area") or ""

        candidatas = []
        for q in questoes:
            area_q = q.get("area") or q.get("grande_area") or q.get("categoria") or ""

            if area_aula and area_q and normalizar(area_aula) != normalizar(area_q):
                continue

            score = calcular_score(aula, q)

            if score <= 0:
                continue

            item = dict(q)
            item["_score_compatibilidade"] = score
            item["_aula_tema"] = aula.get("tema") or aula.get("assunto") or ""
            item["_aula"] = aula.get("aula") or aula.get("numero") or ""
            item["_semana"] = aula.get("semana") or ""

            candidatas.append(item)

        candidatas = sorted(
            candidatas,
            key=lambda x: (
                x.get("_score_compatibilidade", 0),
                x.get("confianca_indexacao") or 0
            ),
            reverse=True
        )

        selecionadas.extend(candidatas[:n_por_area])

    return selecionadas

def texto_enunciado(q):
    for campo in ["enunciado", "texto", "questao"]:
        if q.get(campo):
            return str(q.get(campo)).strip()
    return ""

def alternativas(q):
    itens = []
    for letra in ["a", "b", "c", "d", "e"]:
        for campo in [f"alternativa_{letra}", letra.upper(), letra]:
            if q.get(campo):
                valor = str(q.get(campo)).strip()
                if valor:
                    itens.append((letra.upper(), valor))
                break
    return itens

def gerar_word(questoes, nome_saida=None):
    doc = Document()

    doc.add_heading("ATHENA Question Bank", level=0)
    doc.add_heading("Simulado por Aula — Sprint 8", level=1)

    doc.add_paragraph(
        "Questões selecionadas por compatibilidade entre plano de aula, tema médico, subtema, especialidade e competência."
    )

    for i, q in enumerate(questoes, start=1):
        area = limpar(q.get("area") or q.get("grande_area") or q.get("categoria") or "")
        especialidade = limpar(q.get("especialidade"))
        tema = limpar(q.get("tema_indexado") or q.get("assunto"))
        subtema = limpar(q.get("subtema_indexado"))

        doc.add_heading(f"Questão {i}", level=2)

        p = doc.add_paragraph()
        p.add_run("Área: ").bold = True
        p.add_run(area)

        p = doc.add_paragraph()
        p.add_run("Especialidade: ").bold = True
        p.add_run(especialidade)

        p = doc.add_paragraph()
        p.add_run("Tema: ").bold = True
        p.add_run(tema)

        p = doc.add_paragraph()
        p.add_run("Subtema: ").bold = True
        p.add_run(subtema)

        enun = texto_enunciado(q)
        doc.add_paragraph(enun)

        for letra, alt in alternativas(q):
            doc.add_paragraph(f"{letra}) {alt}")

    if nome_saida is None:
        nome_saida = f"simulado_por_aula_sprint8_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

    caminho = OUTDIR / nome_saida
    doc.save(caminho)
    return caminho

def gerar(aula_numero=None, semana=None, n_por_area=2):
    con = sqlite3.connect(DB)
    con.row_factory = sqlite3.Row
    cur = con.cursor()

    aulas = carregar_aulas(cur, aula_numero=aula_numero, semana=semana)
    questoes = carregar_questoes(cur)

    selecionadas = selecionar_questoes(aulas, questoes, n_por_area=n_por_area)

    caminho = gerar_word(selecionadas)

    con.close()

    print(f"Aulas encontradas: {len(aulas)}")
    print(f"Questões selecionadas: {len(selecionadas)}")
    print(f"Arquivo Word gerado: {caminho}")

    return caminho

if __name__ == "__main__":
    gerar()
