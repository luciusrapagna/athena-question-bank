import argparse
import sqlite3
import csv
from pathlib import Path
from docx import Document

DB_PATH = "app/db/planos_aula.db"
OUTPUT_DIR = Path("outputs/provas_por_aula")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

def buscar_planos(cur, data=None, termo=None):
    sql = """
        SELECT id, data_aula, aula, disciplina, modulo, conteudo
        FROM planos_aula
        WHERE 1=1
    """
    params = []

    if data:
        sql += " AND data_aula = ?"
        params.append(data)

    if termo:
        sql += " AND (LOWER(aula) LIKE ? OR LOWER(conteudo) LIKE ? OR LOWER(disciplina) LIKE ?)"
        busca = f"%{termo.lower()}%"
        params.extend([busca, busca, busca])

    sql += " ORDER BY COALESCE(data_aula, '9999-99-99'), id"

    return cur.execute(sql, params).fetchall()

def buscar_questoes(cur, plano_id, limite):
    return cur.execute("""
        SELECT
            q.id,
            q.grande_area,
            q.tema,
            q.enunciado,
            q.alternativa_a,
            q.alternativa_b,
            q.alternativa_c,
            q.alternativa_d,
            q.alternativa_e,
            q.gabarito,
            ROUND(c.score, 3) AS score
        FROM compatibilidade_plano_questao c
        JOIN questoes q ON q.id = c.questao_id
        WHERE c.plano_id = ?
        ORDER BY c.score DESC
        LIMIT ?
    """, (plano_id, limite)).fetchall()

def exportar_word(plano, questoes):
    plano_id, data_aula, aula, disciplina, modulo, conteudo = plano

    doc = Document()
    doc.add_heading("ATHENA QUESTION BANK", level=1)
    doc.add_heading("Questões sugeridas por plano de aula", level=2)

    doc.add_paragraph(f"Plano ID: {plano_id}")
    doc.add_paragraph(f"Data da aula: {data_aula or 'Não informada'}")
    doc.add_paragraph(f"Disciplina: {disciplina or ''}")
    doc.add_paragraph(f"Módulo: {modulo or ''}")
    doc.add_paragraph(f"Aula: {aula or ''}")
    doc.add_paragraph(f"Total de questões sugeridas: {len(questoes)}")

    doc.add_heading("Questões", level=2)

    for i, q in enumerate(questoes, start=1):
        qid, area, tema, enunciado, a, b, c, d, e, gabarito, score = q
        doc.add_heading(f"Questão {i}", level=3)
        doc.add_paragraph(f"ID: {qid} | Área: {area or ''} | Tema: {tema or ''} | Compatibilidade: {score}")
        doc.add_paragraph(enunciado or "")

        for letra, alt in [("A", a), ("B", b), ("C", c), ("D", d), ("E", e)]:
            if alt:
                doc.add_paragraph(f"{letra}) {alt}")

        if gabarito:
            doc.add_paragraph(f"Gabarito: {gabarito}")

    nome = f"sugestoes_plano_{plano_id}.docx"
    caminho = OUTPUT_DIR / nome
    doc.save(caminho)
    return caminho

def exportar_csv(plano, questoes):
    plano_id = plano[0]
    caminho = OUTPUT_DIR / f"sugestoes_plano_{plano_id}.csv"

    with open(caminho, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f, delimiter=";")
        writer.writerow(["ordem", "questao_id", "grande_area", "tema", "score", "gabarito"])

        for i, q in enumerate(questoes, start=1):
            writer.writerow([i, q[0], q[1], q[2], q[10], q[9] or ""])

    return caminho

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--data", default=None, help="Data da aula no formato salvo no banco")
    parser.add_argument("--termo", default=None, help="Termo para buscar aula/conteúdo/disciplina")
    parser.add_argument("--limite", type=int, default=10)
    parser.add_argument("--listar", action="store_true")
    args = parser.parse_args()

    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()

    planos = buscar_planos(cur, data=args.data, termo=args.termo)

    if args.listar:
        print("\nPLANOS ENCONTRADOS:")
        for p in planos:
            print(f"ID {p[0]} | Data: {p[1]} | Disciplina: {p[3]} | Aula: {p[2]}")
        con.close()
        return

    if not planos:
        print("Nenhum plano encontrado.")
        con.close()
        return

    for plano in planos:
        questoes = buscar_questoes(cur, plano[0], args.limite)

        if not questoes:
            print(f"Plano {plano[0]} sem questões compatíveis.")
            continue

        word = exportar_word(plano, questoes)
        csv_path = exportar_csv(plano, questoes)

        print(f"\nPlano {plano[0]} — {plano[2]}")
        print(f"Word: {word}")
        print(f"CSV: {csv_path}")
        print(f"Questões: {len(questoes)}")

    con.close()

if __name__ == "__main__":
    main()
