import argparse
import sqlite3
from pathlib import Path
from docx import Document

DB_PATH = "app/db/planos_aula.db"
OUTPUT_DIR = Path("outputs/calendario_academico")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

def conectar():
    return sqlite3.connect(DB_PATH)

def listar_sem_data():
    con = conectar()
    cur = con.cursor()

    rows = cur.execute("""
        SELECT id, disciplina, modulo, aula
        FROM planos_aula
        WHERE data_aula IS NULL OR TRIM(data_aula) = ''
        ORDER BY id
    """).fetchall()

    print("\nPLANOS SEM DATA:")
    for r in rows:
        print(f"ID {r[0]} | {r[1] or ''} | {r[2] or ''} | {r[3] or ''}")

    print(f"\nTotal sem data: {len(rows)}")
    con.close()

def definir_data(plano_id, data):
    con = conectar()
    cur = con.cursor()

    cur.execute("""
        UPDATE planos_aula
        SET data_aula = ?
        WHERE id = ?
    """, (data, plano_id))

    con.commit()
    con.close()

    print(f"Data atualizada: plano {plano_id} -> {data}")

def proximas_aulas(data_inicio, data_fim, limite):
    con = conectar()
    cur = con.cursor()

    planos = cur.execute("""
        SELECT id, data_aula, disciplina, modulo, aula, conteudo
        FROM planos_aula
        WHERE data_aula BETWEEN ? AND ?
        ORDER BY data_aula, id
    """, (data_inicio, data_fim)).fetchall()

    if not planos:
        print("Nenhuma aula encontrada no período.")
        con.close()
        return

    doc = Document()
    doc.add_heading("ATHENA QUESTION BANK", level=1)
    doc.add_heading("Calendário Acadêmico Inteligente", level=2)
    doc.add_paragraph(f"Período: {data_inicio} a {data_fim}")

    for plano in planos:
        plano_id, data_aula, disciplina, modulo, aula, conteudo = plano

        questoes = cur.execute("""
            SELECT q.id, q.grande_area, q.tema, q.enunciado, ROUND(c.score, 3)
            FROM compatibilidade_plano_questao c
            JOIN questoes q ON q.id = c.questao_id
            WHERE c.plano_id = ?
            ORDER BY c.score DESC
            LIMIT ?
        """, (plano_id, limite)).fetchall()

        print(f"\n{data_aula} | Plano {plano_id} | {disciplina or ''} | {aula or ''}")
        print(f"Questões sugeridas: {len(questoes)}")

        doc.add_heading(f"{data_aula} — {aula or 'Aula sem título'}", level=2)
        doc.add_paragraph(f"Plano ID: {plano_id}")
        doc.add_paragraph(f"Disciplina: {disciplina or ''}")
        doc.add_paragraph(f"Módulo: {modulo or ''}")
        doc.add_paragraph(f"Questões sugeridas: {len(questoes)}")

        for i, q in enumerate(questoes, start=1):
            qid, area, tema, enunciado, score = q
            print(f"  Q{qid} | {area or ''} | {tema or ''} | score {score}")

            doc.add_paragraph(
                f"Questão {i} — ID {qid} | {area or ''} | {tema or ''} | Compatibilidade: {score}"
            )
            doc.add_paragraph(enunciado or "")

    caminho = OUTPUT_DIR / f"calendario_{data_inicio}_a_{data_fim}.docx"
    doc.save(caminho)

    print(f"\nRelatório gerado: {caminho}")
    con.close()

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--sem-data", action="store_true")
    parser.add_argument("--set-data", nargs=2, metavar=("PLANO_ID", "DATA"))
    parser.add_argument("--inicio", default=None)
    parser.add_argument("--fim", default=None)
    parser.add_argument("--limite", type=int, default=5)

    args = parser.parse_args()

    if args.sem_data:
        listar_sem_data()
        return

    if args.set_data:
        plano_id, data = args.set_data
        definir_data(int(plano_id), data)
        return

    if args.inicio and args.fim:
        proximas_aulas(args.inicio, args.fim, args.limite)
        return

    print("""
Uso:

Listar planos sem data:
python -m app.lesson_planner.calendario_inteligente --sem-data

Definir data de uma aula:
python -m app.lesson_planner.calendario_inteligente --set-data 5 2026-06-10

Gerar sugestões por período:
python -m app.lesson_planner.calendario_inteligente --inicio 2026-06-10 --fim 2026-06-17 --limite 5
""")

if __name__ == "__main__":
    main()
