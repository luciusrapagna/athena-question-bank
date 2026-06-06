import argparse
import sqlite3
import re
import csv
import sys
from pathlib import Path
from difflib import SequenceMatcher
from docx import Document

ROOT_DIR = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(ROOT_DIR))

DB_PATH = "app/db/planos_aula.db"
ROOT_DIR = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(ROOT_DIR))

from app.medical_dictionary.expansor_medico import expandir_termo_medico

OUTPUT_DIR = Path("outputs/simulados_por_aula")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

AREAS = [
    "Clínica Médica",
    "Cirurgia",
    "Pediatria",
    "Ginecologia e Obstetrícia",
    "Saúde Coletiva",
]

def normalizar(texto):
    import unicodedata
    texto = str(texto or "").lower()
    texto = unicodedata.normalize("NFKD", texto)
    texto = "".join(c for c in texto if not unicodedata.combining(c))
    texto = re.sub(r"[^a-z0-9\s]", " ", texto)
    return re.sub(r"\s+", " ", texto).strip()


STOPWORDS = {
    "de", "da", "do", "das", "dos", "e", "em", "na", "no", "nas", "nos",
    "para", "por", "com", "sem", "ao", "aos", "as", "os", "um", "uma",
    "aplicar", "analisar", "reconhecer", "identificar", "propor",
    "paciente", "pacientes", "clinica", "criterios", "tratamento"
}


EXPANSOES_MEDICAS = {
    "angina": "dor toracica cardiovascular coronariana isquemia miocardio infarto sindrome coronariana",
    "has": "hipertensao arterial risco cardiovascular pressao arterial",
    "hipertensao": "pressao arterial has risco cardiovascular",
    "cardiologico": "cardiologia cardiovascular coracao ausculta semiologia exame fisico",
    "pre natal": "gestacao gestante obstetricia gravidez acompanhamento prenatal",
    "sepse": "infeccao choque septico antibiotico emergencia uti",
    "diabetes": "glicemia insulinoterapia endocrinologia hiperglicemia",
}


def expandir_texto(texto):
    base = normalizar(texto)
    texto_expandido = expandir_termo_medico(base, usar_api=False)

    acrescimos = []
    for chave, expansao in EXPANSOES_MEDICAS.items():
        if chave in base:
            acrescimos.append(expansao)

    return normalizar(texto_expandido + " " + " ".join(acrescimos))


def tokens_relevantes(texto):
    toks = expandir_texto(texto).split()
    return {t for t in toks if len(t) >= 4 and t not in STOPWORDS}


def score_texto(a, b):
    a_norm = expandir_texto(a)
    b_norm = expandir_texto(b)

    if not a_norm or not b_norm:
        return 0

    termos_a = tokens_relevantes(a_norm)
    termos_b = tokens_relevantes(b_norm)

    if not termos_a or not termos_b:
        return 0

    inter = termos_a & termos_b
    uniao = termos_a | termos_b

    jaccard = len(inter) / max(len(uniao), 1)
    cobertura_aula = len(inter) / max(len(termos_a), 1)
    cobertura_questao = len(inter) / max(len(termos_b), 1)
    seq = SequenceMatcher(None, a_norm, b_norm).ratio()

    score = (
        0.45 * cobertura_aula +
        0.25 * jaccard +
        0.20 * cobertura_questao +
        0.10 * seq
    )

    return round(score, 4)


def buscar_aula(cur, periodo=None, aula_numero=None, data=None, termo=None):
    sql = """
        SELECT id, plano_id, aula_numero, data_aula, tema, objetivos
        FROM aulas_cronograma
        WHERE aula_numero GLOB '[0-9]*'
    """
    params = []

    if periodo:
        sql += " AND plano_id = ?"
        params.append(periodo)

    if aula_numero:
        sql += " AND CAST(aula_numero AS INTEGER) = ?"
        params.append(aula_numero)

    if data:
        sql += " AND data_aula = ?"
        params.append(data)

    if termo:
        sql += " AND (LOWER(tema) LIKE ? OR LOWER(objetivos) LIKE ?)"
        busca = f"%{termo.lower()}%"
        params.extend([busca, busca])

    sql += " ORDER BY plano_id, CAST(aula_numero AS INTEGER), data_aula LIMIT 1"
    return cur.execute(sql, params).fetchone()

def buscar_questoes_por_area(cur, area):
    return cur.execute("""
        SELECT id, grande_area, tema, enunciado,
               alternativa_a, alternativa_b, alternativa_c, alternativa_d, alternativa_e,
               gabarito, fonte, ano, prova, instituicao,
               subarea, comentario, competencia, habilidade
        FROM questoes
        WHERE grande_area = ?
          AND enunciado IS NOT NULL
          AND TRIM(enunciado) <> ''
          AND tema IS NOT NULL
          AND TRIM(tema) <> ''
          AND LOWER(TRIM(tema)) NOT IN (
              'assunto não identificado',
              'tema não identificado',
              'não identificado',
              'nao identificado',
              'sem tema'
          )
          AND grande_area IS NOT NULL
          AND TRIM(grande_area) <> ''
          AND LOWER(TRIM(grande_area)) NOT IN (
              'área não identificada',
              'area não identificada',
              'area nao identificada',
              'não identificado',
              'nao identificado'
          )
    """, (area,)).fetchall()

def aula_eh_feriado(aula):
    texto = " ".join(str(x or "") for x in aula).lower()
    return "feriado" in texto


def resolver_aula_efetiva(cur, aula):
    """
    Se a aula selecionada for FERIADO, usa o conteúdo da aula anterior válida
    do mesmo plano de aula.
    """
    if not aula_eh_feriado(aula):
        return aula

    aula_id, plano_id, aula_numero, data_aula, tema, objetivos = aula

    anterior = cur.execute("""
        SELECT id, plano_id, aula_numero, data_aula, tema, objetivos
        FROM aulas_cronograma
        WHERE plano_id = ?
          AND CAST(aula_numero AS INTEGER) < CAST(? AS INTEGER)
          AND LOWER(COALESCE(tema, '')) NOT LIKE '%feriado%'
          AND LOWER(COALESCE(objetivos, '')) NOT LIKE '%feriado%'
        ORDER BY CAST(aula_numero AS INTEGER) DESC
        LIMIT 1
    """, (plano_id, aula_numero)).fetchone()

    return anterior or aula


def ranquear(aula, questoes, score_minimo=0.07):
    _, _, _, _, tema_aula, objetivos = aula

    texto_aula = f"{tema_aula or ''} {tema_aula or ''} {objetivos or ''}"
    resultado = []

    for q in questoes:
        tema_q = q[2] if len(q) > 2 else ""
        enunciado = q[3] if len(q) > 3 else ""
        subarea = q[14] if len(q) > 14 else ""
        comentario = q[15] if len(q) > 15 else ""
        competencia = q[16] if len(q) > 16 else ""
        habilidade = q[17] if len(q) > 17 else ""

        texto_q = " ".join([
            str(tema_q or ""),
            str(tema_q or ""),
            str(tema_q or ""),
            str(subarea or ""),
            str(enunciado or ""),
            str(enunciado or ""),
            str(comentario or ""),
            str(competencia or ""),
            str(habilidade or ""),
        ])

        score = score_texto(texto_aula, texto_q)

        if score >= score_minimo:
            resultado.append((score, q))

    resultado.sort(key=lambda x: x[0], reverse=True)
    return resultado


def exportar_word(aula, selecionadas, nome):
    aula_id, plano_id, aula_numero, data_aula, tema, objetivos = aula

    doc = Document()
    doc.add_heading("ATHENA QUESTION BANK", level=1)
    doc.add_heading("Simulado por Aula", level=2)

    doc.add_paragraph(f"Aula: {aula_numero}")
    doc.add_paragraph(f"Data: {data_aula}")
    doc.add_paragraph(f"Tema: {tema}")
    doc.add_paragraph(f"Objetivos: {objetivos}")

    total = sum(len(v) for v in selecionadas.values())
    doc.add_paragraph(f"Total de questões: {total}")

    n = 1
    for area, itens in selecionadas.items():
        doc.add_heading(area, level=2)

        for score, q in itens:
            qid, grande_area, tema_q, enunciado, a, b, c, d, e, gabarito, fonte, ano, prova, instituicao, subarea, comentario, competencia, habilidade = q

            doc.add_heading(f"Questão {n} — ID {qid}", level=3)
            doc.add_paragraph(f"Compatibilidade: {score}")
            doc.add_paragraph(f"Área: {grande_area or ''} | Tema: {tema_q or ''}")
            doc.add_paragraph(f"Fonte: {fonte or ''} | Prova: {prova or ''} | Ano: {ano or ''} | Instituição: {instituicao or ''}")
            doc.add_paragraph(enunciado or "")

            for letra, alt in [("A", a), ("B", b), ("C", c), ("D", d), ("E", e)]:
                if alt:
                    doc.add_paragraph(f"{letra}) {alt}")

            n += 1

    caminho = OUTPUT_DIR / nome
    doc.save(caminho)
    return caminho

def exportar_csv(aula, selecionadas, nome):
    caminho = OUTPUT_DIR / nome

    with open(caminho, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f, delimiter=";")
        writer.writerow([
            "numero", "aula_numero", "data_aula", "tema_aula",
            "questao_id", "score", "grande_area", "tema_questao",
            "fonte", "prova", "ano", "instituicao", "gabarito"
        ])

        n = 1
        for area, itens in selecionadas.items():
            for score, q in itens:
                qid, grande_area, tema_q, enunciado, a, b, c, d, e, gabarito, fonte, ano, prova, instituicao, subarea, comentario, competencia, habilidade = q
                writer.writerow([
                    n, aula[2], aula[3], aula[4],
                    qid, score, grande_area, tema_q,
                    fonte, prova, ano, instituicao, gabarito
                ])
                n += 1

    return caminho

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--periodo", type=int, default=None)
    parser.add_argument("--aula", type=int, default=None)
    parser.add_argument("--data", default=None)
    parser.add_argument("--termo", default=None)
    parser.add_argument("--clinica", type=int, default=0)
    parser.add_argument("--cirurgia", type=int, default=0)
    parser.add_argument("--pediatria", type=int, default=0)
    parser.add_argument("--go", type=int, default=0)
    parser.add_argument("--saude-coletiva", type=int, default=0)
    args = parser.parse_args()

    quotas = {
        "Clínica Médica": args.clinica,
        "Cirurgia": args.cirurgia,
        "Pediatria": args.pediatria,
        "Ginecologia e Obstetrícia": args.go,
        "Saúde Coletiva": args.saude_coletiva,
    }

    quotas = {area: qtd for area, qtd in quotas.items() if qtd > 0}

    if not quotas:
        print("Informe pelo menos uma quantidade por área.")
        return

    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()

    aula = buscar_aula(cur, periodo=args.periodo, aula_numero=args.aula, data=args.data, termo=args.termo)

    if not aula:
        print("Nenhuma aula encontrada.")
        con.close()
        return

    aula_original = aula
    aula = resolver_aula_efetiva(cur, aula_original)

    selecionadas = {}

    print(f"\nAula solicitada: {aula_original[2]} | {aula_original[3]} | {aula_original[4]}")

    if aula_original != aula:
        print(f"Aula solicitada é FERIADO. Usando conteúdo da aula anterior válida:")
        print(f"Aula efetiva: {aula[2]} | {aula[3]} | {aula[4]}")
    else:
        print(f"Aula efetiva: {aula[2]} | {aula[3]} | {aula[4]}")

    print("Objetivos:", aula[5])

    for area, qtd in quotas.items():
        questoes = buscar_questoes_por_area(cur, area)
        ranqueadas = ranquear(aula, questoes)
        selecionadas[area] = ranqueadas[:qtd]

        print(f"\n{area}: solicitadas {qtd}, selecionadas {len(selecionadas[area])}")

        if len(selecionadas[area]) == 0:
            print("  Nenhuma questão compatível encontrada com o tema/objetivos da aula.")
            print("  Sugestão: revisar classificação temática do banco ou reduzir temporariamente o score mínimo.")
        else:
            for score, q in selecionadas[area]:
                print(f"  Q{q[0]} | score {score} | {q[2]} | Fonte: {q[10]} | Ano: {q[11]}")

    if args.periodo and args.aula:
        sufixo = f"P{args.periodo:02d}_A{args.aula:02d}"
    else:
        sufixo = args.data or normalizar(args.termo or "aula").replace(" ", "_")
    word = exportar_word(aula, selecionadas, f"simulado_aula_{sufixo}.docx")
    csv_path = exportar_csv(aula, selecionadas, f"simulado_aula_{sufixo}.csv")

    print(f"\nWord: {word}")
    print(f"CSV: {csv_path}")

    con.close()

if __name__ == "__main__":
    main()
