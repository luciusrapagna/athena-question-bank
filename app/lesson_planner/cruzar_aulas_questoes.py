import argparse
import sqlite3
import re
from pathlib import Path
from docx import Document
import csv
from difflib import SequenceMatcher

DB_PATH = "app/db/planos_aula.db"
OUTPUT_DIR = Path("outputs/aulas_questoes")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

def normalizar(texto):
    texto = (texto or "").lower()
    texto = re.sub(r"[^a-zà-ÿ0-9\s]", " ", texto)
    texto = re.sub(r"\s+", " ", texto)
    return texto.strip()

def score_texto(a, b):
    a = normalizar(a)
    b = normalizar(b)
    if not a or not b:
        return 0
    termos_a = set(a.split())
    termos_b = set(b.split())
    inter = len(termos_a & termos_b)
    jaccard = inter / max(len(termos_a | termos_b), 1)
    seq = SequenceMatcher(None, a, b).ratio()
    return round((0.7 * jaccard) + (0.3 * seq), 4)

def buscar_aulas(cur, data=None, termo=None):
    sql = """
        SELECT id, plano_id, aula_numero, data_aula, tema, objetivos
        FROM aulas_cronograma
        WHERE 1=1
    """
    params = []

    if data:
        sql += " AND data_aula = ?"
        params.append(data)

    if termo:
        sql += " AND (LOWER(tema) LIKE ? OR LOWER(objetivos) LIKE ?)"
        busca = f"%{termo.lower()}%"
        params.extend([busca, busca])

    sql += " ORDER BY data_aula, plano_id, CAST(aula_numero AS INTEGER)"
    return cur.execute(sql, params).fetchall()

def buscar_questoes(cur):
    return cur.execute("""
        SELECT id, grande_area, tema, enunciado,
               alternativa_a, alternativa_b, alternativa_c, alternativa_d, alternativa_e,
               gabarito, fonte, ano, prova, instituicao
        FROM questoes
        WHERE enunciado IS NOT NULL
          AND TRIM(enunciado) <> ''
    """).fetchall()

def ranquear_questoes(aula, questoes, limite):
    _, _, _, _, tema_aula, objetivos = aula
    texto_aula = f"{tema_aula or ''} {objetivos or ''}"

    ranqueadas = []

    for q in questoes:
        qid, area, tema_q, enunciado, a, b, c, d, e, gabarito, fonte, ano, prova, instituicao = q
        texto_q = f"{area or ''} {tema_q or ''} {enunciado or ''} {a or ''} {b or ''} {c or ''} {d or ''} {e or ''}"
        score = score_texto(texto_aula, texto_q)

        if score > 0:
            ranqueadas.append((score, q))

    ranqueadas.sort(key=lambda x: x[0], reverse=True)
    return ranqueadas[:limite]

def exportar_word(aulas_resultado, nome):
    doc = Document()
    doc.add_heading("ATHENA QUESTION BANK", level=1)
    doc.add_heading("Questões sugeridas por aula do cronograma", level=2)

    for aula, sugestoes in aulas_resultado:
        aula_id, plano_id, aula_numero, data_aula, tema, objetivos = aula

        doc.add_heading(f"Aula {aula_numero} — {data_aula or 'Sem data'}", level=2)
        doc.add_paragraph(f"Plano ID: {plano_id}")
        doc.add_paragraph(f"Tema: {tema or ''}")
        doc.add_paragraph(f"Objetivos: {objetivos or ''}")
        doc.add_paragraph(f"Questões sugeridas: {len(sugestoes)}")

        for i, (score, q) in enumerate(sugestoes, start=1):
            qid, area, tema_q, enunciado, a, b, c, d, e, gabarito, fonte, ano, prova, instituicao = q

            doc.add_heading(f"Questão sugerida {i} — ID {qid}", level=3)
            doc.add_paragraph(f"Compatibilidade: {score}")
            doc.add_paragraph(f"Área: {area or ''} | Tema da questão: {tema_q or ''}")
            doc.add_paragraph(f"Fonte: {fonte or ''} | Prova: {prova or ''} | Ano: {ano or ''} | Instituição: {instituicao or ''}")
            doc.add_paragraph(enunciado or "")

            for letra, alt in [("A", a), ("B", b), ("C", c), ("D", d), ("E", e)]:
                if alt:
                    doc.add_paragraph(f"{letra}) {alt}")

            if gabarito:
                doc.add_paragraph(f"Gabarito: {gabarito}")

    caminho = OUTPUT_DIR / nome
    doc.save(caminho)
    return caminho

def exportar_csv(aulas_resultado, nome):
    caminho = OUTPUT_DIR / nome

    with open(caminho, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f, delimiter=";")
        writer.writerow([
            "aula_id", "plano_id", "aula_numero", "data_aula", "tema_aula",
            "questao_id", "score", "grande_area", "tema_questao",
            "fonte", "prova", "ano", "instituicao", "gabarito"
        ])

        for aula, sugestoes in aulas_resultado:
            aula_id, plano_id, aula_numero, data_aula, tema, objetivos = aula

            for score, q in sugestoes:
                qid, area, tema_q, enunciado, a, b, c, d, e, gabarito, fonte, ano, prova, instituicao = q

                writer.writerow([
                    aula_id, plano_id, aula_numero, data_aula, tema,
                    qid, score, area, tema_q, fonte, prova, ano, instituicao, gabarito
                ])

    return caminho

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--data", default=None)
    parser.add_argument("--termo", default=None)
    parser.add_argument("--limite", type=int, default=5)
    args = parser.parse_args()

    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()

    aulas = buscar_aulas(cur, data=args.data, termo=args.termo)
    questoes = buscar_questoes(cur)

    if not aulas:
        print("Nenhuma aula encontrada em aulas_cronograma.")
        print("Execute antes: python -m app.lesson_planner.extrair_cronograma_planos")
        con.close()
        return

    aulas_resultado = []

    for aula in aulas:
        sugestoes = ranquear_questoes(aula, questoes, args.limite)
        aulas_resultado.append((aula, sugestoes))

        print(f"\nAula {aula[2]} | {aula[3]} | {aula[4]}")
        for score, q in sugestoes:
            print(f"  Q{q[0]} | score {score} | {q[1]} | {q[2]} | Fonte: {q[10]} | Ano: {q[11]}")

    sufixo = args.data or args.termo or "geral"
    word = exportar_word(aulas_resultado, f"aulas_questoes_{sufixo}.docx")
    csv_path = exportar_csv(aulas_resultado, f"aulas_questoes_{sufixo}.csv")

    print(f"\nWord: {word}")
    print(f"CSV: {csv_path}")

    con.close()

if __name__ == "__main__":
    main()
