import sqlite3
import csv
from pathlib import Path
from docx import Document

DB_PATH = "app/db/planos_aula.db"
OUTPUT_DIR = Path("outputs/simulados")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

def conectar():
    return sqlite3.connect(DB_PATH)

def listar_areas(cur):
    return cur.execute("""
        SELECT grande_area, COUNT(*)
        FROM questoes
        WHERE grande_area IS NOT NULL AND grande_area <> ''
        GROUP BY grande_area
        ORDER BY grande_area
    """).fetchall()

def listar_temas(cur, area):
    return cur.execute("""
        SELECT tema, COUNT(*)
        FROM questoes
        WHERE grande_area = ?
          AND tema IS NOT NULL
          AND tema <> ''
        GROUP BY tema
        ORDER BY tema
    """, (area,)).fetchall()

def buscar_questoes(cur, area, temas, limite):
    placeholders = ",".join("?" for _ in temas)
    sql = f"""
        SELECT id, grande_area, tema, enunciado,
               alternativa_a, alternativa_b, alternativa_c,
               alternativa_d, alternativa_e, gabarito, fonte, ano
        FROM questoes
        WHERE grande_area = ?
          AND tema IN ({placeholders})
        ORDER BY RANDOM()
        LIMIT ?
    """
    return cur.execute(sql, [area] + temas + [limite]).fetchall()

def escolher_indices(texto, maximo):
    if not texto.strip():
        return list(range(1, maximo + 1))
    indices = []
    for item in texto.split(","):
        item = item.strip()
        if item.isdigit():
            n = int(item)
            if 1 <= n <= maximo:
                indices.append(n)
    return indices

def exportar_word(questoes, nome="simulado_inteligente.docx"):
    doc = Document()
    doc.add_heading("ATHENA QUESTION BANK", level=1)
    doc.add_heading("Simulado Inteligente", level=2)
    doc.add_paragraph(f"Total de questões: {len(questoes)}")

    for i, q in enumerate(questoes, start=1):
        qid, area, tema, enunciado, a, b, c, d, e, gabarito, fonte, ano = q
        doc.add_heading(f"Questão {i}", level=2)
        doc.add_paragraph(f"ID: {qid} | Área: {area} | Tema: {tema}")
        doc.add_paragraph(f"Fonte: {fonte or ''} | Ano: {ano or ''}")
        doc.add_paragraph(enunciado or "")

        for letra, alt in [("A", a), ("B", b), ("C", c), ("D", d), ("E", e)]:
            if alt:
                doc.add_paragraph(f"{letra}) {alt}")

    caminho = OUTPUT_DIR / nome
    doc.save(caminho)
    return caminho

def exportar_gabarito(questoes, nome="gabarito_simulado_inteligente.csv"):
    caminho = OUTPUT_DIR / nome
    with open(caminho, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f, delimiter=";")
        writer.writerow(["numero", "questao_id", "grande_area", "tema", "gabarito"])
        for i, q in enumerate(questoes, start=1):
            writer.writerow([i, q[0], q[1], q[2], q[9] or ""])
    return caminho

def main():
    con = conectar()
    cur = con.cursor()

    selecionadas = []

    print("\nATHENA SIMULADO INTELIGENTE")
    print("=" * 60)

    areas = listar_areas(cur)

    for i, (area, total) in enumerate(areas, start=1):
        print(f"{i}. {area} ({total} questões)")

    escolha_areas = input("\nEscolha as áreas por número, separadas por vírgula, ou ENTER para todas: ")
    indices_areas = escolher_indices(escolha_areas, len(areas))

    for idx in indices_areas:
        area = areas[idx - 1][0]
        temas = listar_temas(cur, area)

        if not temas:
            print(f"\n{area}: sem temas disponíveis.")
            continue

        print(f"\n{area}")
        print("-" * 60)

        for j, (tema, total) in enumerate(temas, start=1):
            print(f"{j}. {tema} ({total} questões)")

        escolha_temas = input("Escolha os temas por número, separados por vírgula, ou ENTER para todos: ")
        indices_temas = escolher_indices(escolha_temas, len(temas))
        temas_escolhidos = [temas[j - 1][0] for j in indices_temas]

        qtd_txt = input(f"Quantas questões deseja para {area}? ").strip()
        if not qtd_txt.isdigit():
            print("Quantidade inválida. Área ignorada.")
            continue

        qtd = int(qtd_txt)
        questoes = buscar_questoes(cur, area, temas_escolhidos, qtd)
        selecionadas.extend(questoes)

        print(f"{len(questoes)} questões adicionadas de {area}.")

    con.close()

    if not selecionadas:
        print("\nNenhuma questão selecionada.")
        return

    word = exportar_word(selecionadas)
    gabarito = exportar_gabarito(selecionadas)

    print("\nSimulado gerado com sucesso.")
    print(f"Word: {word}")
    print(f"Gabarito: {gabarito}")
    print(f"Total: {len(selecionadas)} questões")

if __name__ == "__main__":
    main()
