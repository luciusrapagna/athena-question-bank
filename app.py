import sqlite3
from datetime import datetime
from io import BytesIO
from pathlib import Path

import pandas as pd
import streamlit as st
from docx import Document

try:
    from app.config_paths import BANCO_QUESTOES, BANCO_PLANOS, BANCO_RELACIONAMENTOS
except Exception:
    BASE_TMP = Path(__file__).resolve().parent
    BANCO_QUESTOES = BASE_TMP / "data" / "banco_questoes" / "banco_questoes.csv"
    BANCO_PLANOS = BASE_TMP / "data" / "banco_planos" / "banco_planos.csv"
    BANCO_RELACIONAMENTOS = BASE_TMP / "data" / "banco_relacionamentos" / "questoes_planos.csv"

BASE = Path(__file__).resolve().parent
DATA = BASE / "data"
PROVAS = DATA / "entrada" / "provas"
PLANOS = DATA / "entrada" / "planos_aula"
DB = DATA / "athena_question_bank.db"
OUT_WORD = BASE / "outputs" / "word"
OUT_REL = BASE / "outputs" / "relatorios"
OUT_SIM = BASE / "outputs" / "simulados"

for p in [DATA, PROVAS, PLANOS, OUT_WORD, OUT_REL, OUT_SIM]:
    p.mkdir(parents=True, exist_ok=True)

def init_db():
    con = sqlite3.connect(DB)
    cur = con.cursor()
    cur.execute("""CREATE TABLE IF NOT EXISTS arquivos (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        tipo TEXT,
        nome TEXT,
        caminho TEXT,
        extensao TEXT,
        tamanho INTEGER,
        data_upload TEXT
    )""")
    con.commit()
    con.close()

def registrar(tipo, caminho):
    con = sqlite3.connect(DB)
    cur = con.cursor()
    cur.execute(
        "INSERT INTO arquivos (tipo,nome,caminho,extensao,tamanho,data_upload) VALUES (?,?,?,?,?,?)",
        (tipo, caminho.name, str(caminho), caminho.suffix.lower(), caminho.stat().st_size, datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    )
    con.commit()
    con.close()

def tabela_arquivos(tipo):
    con = sqlite3.connect(DB)
    df = pd.read_sql_query("SELECT * FROM arquivos WHERE tipo=? ORDER BY id DESC", con, params=(tipo,))
    con.close()
    return df

def salvar_upload(files, destino, tipo):
    for f in files:
        caminho = destino / f.name
        caminho.write_bytes(f.getbuffer())
        registrar(tipo, caminho)

def ler_csv(path):
    path = Path(path)
    if not path.exists():
        return pd.DataFrame()
    return pd.read_csv(path, encoding="utf-8-sig").fillna("")

def gerar_docx(questoes, titulo):
    doc = Document()
    doc.add_heading("ATHENA QUESTION BANK", level=1)
    doc.add_heading(titulo, level=2)
    doc.add_paragraph(f"Total de questões: {len(questoes)}")
    for i, q in enumerate(questoes, 1):
        doc.add_heading(f"Questão {i}", level=2)
        doc.add_paragraph(f"Área: {q.get('area','')} | Assunto: {q.get('assunto','')} | Competência: {q.get('competencia','')}")
        doc.add_paragraph(str(q.get("enunciado", "")))
        for letra, campo in [("A","alternativa_a"),("B","alternativa_b"),("C","alternativa_c"),("D","alternativa_d"),("E","alternativa_e")]:
            if q.get(campo):
                doc.add_paragraph(f"{letra}) {q.get(campo)}")
        if q.get("gabarito"):
            doc.add_paragraph(f"Gabarito: {q.get('gabarito')}")
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

init_db()

st.set_page_config(page_title="ATHENA Question Bank", page_icon="🧠", layout="wide")
st.title("🧠 ATHENA Question Bank")
st.caption("Upload de provas • Upload de planos • Banco persistente • Question Selector • Exam Analytics • Simulado Generator")

st.subheader("1. Inserção de dados")

c1, c2 = st.columns(2)
with c1:
    provas = st.file_uploader("Inserir provas", type=["pdf","docx","txt","md","csv","xlsx"], accept_multiple_files=True)
    if provas and st.button("Salvar provas", use_container_width=True):
        salvar_upload(provas, PROVAS, "prova")
        st.success("Provas salvas no banco persistente.")

with c2:
    planos = st.file_uploader("Inserir planos de aula", type=["pdf","docx","txt","md","csv","xlsx"], accept_multiple_files=True)
    if planos and st.button("Salvar planos", use_container_width=True):
        salvar_upload(planos, PLANOS, "plano")
        st.success("Planos salvos no banco persistente.")

df_q = ler_csv(BANCO_QUESTOES)
df_p = ler_csv(BANCO_PLANOS)
df_r = ler_csv(BANCO_RELACIONAMENTOS)

m1, m2, m3 = st.columns(3)
m1.metric("Provas armazenadas", len(tabela_arquivos("prova")))
m2.metric("Planos armazenados", len(tabela_arquivos("plano")))
m3.metric("Questões no banco", len(df_q))

with st.expander("📦 Ver banco persistente de arquivos"):
    st.write("Provas")
    st.dataframe(tabela_arquivos("prova"), use_container_width=True)
    st.write("Planos de aula")
    st.dataframe(tabela_arquivos("plano"), use_container_width=True)

st.divider()

aba0, aba_curadoria, aba1, aba2, aba3 = st.tabs(["🔄 Processar Banco", "🧹 Curadoria do Banco", "🎯 Question Selector", "📊 Exam Analytics", "📝 Simulado Generator"])


with aba0:
    st.header("🔄 Processar Banco")
    st.info("Processa uploads de provas e planos, atualiza o banco SQLite e cria relacionamentos Questões × Planos.")

    if st.button("Executar Processamento do Banco", type="primary"):
        try:
            from app.processar_banco.pipeline import processar_tudo
            resultado_proc = processar_tudo()
            st.success("Banco processado com sucesso.")
            st.json(resultado_proc)
            st.cache_data.clear()
        except Exception as e:
            st.error("Erro ao processar banco.")
            st.exception(e)



with aba_curadoria:
    st.header("🧹 Curadoria do Banco")
    st.info("Diagnóstico automático da qualidade do banco de questões.")

    col_diag, col_corr = st.columns(2)

    with col_diag:
        executar_diag = st.button("Executar Diagnóstico de Curadoria", type="primary", use_container_width=True)

    with col_corr:
        executar_corr = st.button("Executar Correção Automática", type="secondary", use_container_width=True)

    if executar_corr:
        try:
            from app.curadoria.diagnostico import corrigir_automaticamente
            resultado_correcao = corrigir_automaticamente()
            st.success("Correção automática executada.")
            st.json(resultado_correcao)
            st.cache_data.clear()
        except Exception as e:
            st.error("Erro ao executar correção automática.")
            st.exception(e)

    if executar_diag or executar_corr:
        try:
            from app.curadoria.diagnostico import diagnosticar_banco
            resumo_curadoria, problemas_curadoria, caminho_relatorio = diagnosticar_banco()

            st.subheader("Resumo da Curadoria")
            st.dataframe(resumo_curadoria, use_container_width=True)

            st.subheader("Questões que precisam de revisão")
            if problemas_curadoria.empty:
                st.success("Nenhum problema encontrado no banco.")
            else:
                st.dataframe(problemas_curadoria, use_container_width=True)

                csv = problemas_curadoria.to_csv(index=False).encode("utf-8-sig")
                st.download_button(
                    "Baixar relatório de curadoria CSV",
                    csv,
                    "curadoria_banco_questoes.csv",
                    "text/csv"
                )

            st.caption(f"Relatório salvo em: {caminho_relatorio}")

        except Exception as e:
            st.error("Erro ao executar curadoria.")
            st.exception(e)

    st.divider()
    st.caption("A correção automática altera apenas campos seguros: área, assunto e competência. Enunciado e alternativas não são modificados.")

with aba1:
    st.header("🎯 Question Selector")
    if df_q.empty:
        st.warning("Banco de questões não encontrado.")
    else:
        filtrado = df_q.copy()

        areas = sorted(filtrado["area"].dropna().astype(str).unique()) if "area" in filtrado.columns else []
        area = st.selectbox("Grande área", [""] + areas)

        if area:
            filtrado = filtrado[filtrado["area"].astype(str) == area]

        assuntos = sorted(filtrado["assunto"].dropna().astype(str).unique()) if "assunto" in filtrado.columns else []
        assunto = st.selectbox("Assunto", [""] + assuntos)

        if assunto:
            filtrado = filtrado[filtrado["assunto"].astype(str) == assunto]

        competencias = sorted(filtrado["competencia"].dropna().astype(str).unique()) if "competencia" in filtrado.columns else []
        competencia = st.selectbox("Competência", [""] + competencias)

        if competencia:
            filtrado = filtrado[filtrado["competencia"].astype(str) == competencia]

        qtd = st.number_input("Quantidade", min_value=1, max_value=max(1, len(filtrado)), value=min(5, max(1, len(filtrado))))
        resultado = filtrado.head(int(qtd))

        st.dataframe(resultado, use_container_width=True)

        st.download_button("Baixar CSV", resultado.to_csv(index=False).encode("utf-8-sig"), "question_selector.csv", "text/csv")
        st.download_button("Baixar Word", gerar_docx(resultado.to_dict("records"), "Question Selector"), "question_selector.docx")

with aba2:
    st.header("📊 Exam Analytics")
    if df_q.empty or "area" not in df_q.columns:
        st.warning("Banco sem coluna de área.")
    else:
        total = len(df_q)
        areas_df = df_q["area"].value_counts().reset_index()
        areas_df.columns = ["grande_area", "n_questoes"]
        areas_df["percentual_total"] = (areas_df["n_questoes"] / total * 100).round(2)

        st.subheader("Percentual das grandes áreas")
        st.dataframe(areas_df, use_container_width=True)
        st.bar_chart(areas_df.set_index("grande_area")["percentual_total"])

        if "assunto" in df_q.columns:
            assuntos_df = df_q.groupby("area")["assunto"].value_counts().reset_index(name="n_questoes")
            assuntos_df["percentual_dentro_area"] = assuntos_df.groupby("area")["n_questoes"].transform(lambda x: (x / x.sum() * 100).round(2))
            st.subheader("Percentual dos assuntos por área")
            st.dataframe(assuntos_df, use_container_width=True)

            blueprint = assuntos_df.merge(areas_df[["grande_area","percentual_total"]], left_on="area", right_on="grande_area", how="left")
            st.subheader("Blueprint automático")
            st.dataframe(blueprint, use_container_width=True)

with aba3:
    st.header("📝 Simulado Generator")
    if df_q.empty:
        st.warning("Banco de questões não encontrado.")
    else:
        sim = df_q.copy()

        areas = sorted(sim["area"].dropna().astype(str).unique()) if "area" in sim.columns else []
        areas_sel = st.multiselect("Grandes áreas", areas)
        if areas_sel:
            sim = sim[sim["area"].astype(str).isin(areas_sel)]

        assuntos = sorted(sim["assunto"].dropna().astype(str).unique()) if "assunto" in sim.columns else []
        assuntos_sel = st.multiselect("Assuntos", assuntos)
        if assuntos_sel:
            sim = sim[sim["assunto"].astype(str).isin(assuntos_sel)]

        qtd = st.number_input("Número de questões do simulado", min_value=1, max_value=max(1, len(sim)), value=min(10, max(1, len(sim))))
        sim_final = sim.sample(n=int(qtd), random_state=42) if len(sim) >= int(qtd) else sim

        st.dataframe(sim_final, use_container_width=True)
        st.download_button("Exportar CSV", sim_final.to_csv(index=False).encode("utf-8-sig"), "simulado.csv", "text/csv")
        st.download_button("Exportar Word", gerar_docx(sim_final.to_dict("records"), "Simulado Generator"), "simulado.docx")

st.divider()
st.caption("ATHENA Question Bank • Integrado ao Ecossistema ATHENA Scientific")
